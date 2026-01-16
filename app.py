from flask import Flask, render_template, request, redirect, send_file
import json
import csv
from datetime import datetime, timedelta, time
import os
import re
from collections import defaultdict
from docx import Document
from docx.shared import Pt

app = Flask(__name__)

# --- File Paths Configuration -------------------------------------------------
# Data files
DATA_DIR = "data"
LOG_FILE = os.path.join(DATA_DIR, "logs.json")
NAMES_FILE = os.path.join(DATA_DIR, "names.json")
RFID_FILE = os.path.join(DATA_DIR, "rfid_cards.json")

# Template files
DOCX_TEMPLATES_DIR = "docx_templates"
TEMPLATE_DOCX = os.path.join(DOCX_TEMPLATES_DIR, "FH306 Sign-In Sheet.docx")

# Export Configuration
EXPORTS_DIR = "exports"
DOCX_EXPORT_FOLDER = os.path.join(EXPORTS_DIR, "docx_exports")
CSV_BASE_FILENAME = os.path.join(EXPORTS_DIR, "room_logs.csv")
CSV_DATE_FILENAME_TEMPLATE = os.path.join(EXPORTS_DIR, "room_logs_{date}.csv")
DOCX_OUTPUT_FILENAME_TEMPLATE = "FH306 Sign-In Sheet - {date}.docx"

# --- Utility Functions -------------------------------------------------------

def load_json(path, default):
    if not os.path.exists(path):
        with open(path, "w") as f:
            json.dump(default, f)
    with open(path, "r") as f:
        return json.load(f)

def save_json(path, data):
    with open(path, "w") as f:
        json.dump(data, f, indent=4)

def load_logs():
    return load_json(LOG_FILE, [])

def save_logs(logs):
    save_json(LOG_FILE, logs)

def remove_seconds_from_logs():
    """
    Update all existing logs to remove seconds from timestamps.
    Converts "2025-11-20 05:01:22 PM" to "2025-11-20 5:01 PM"
    """
    logs = load_logs()
    updated_count = 0
    
    for log in logs:
        timestamp = log.get("timestamp", "")
        if not timestamp:
            continue
        
        # Check if timestamp has seconds (format: "YYYY-MM-DD HH:MM:SS AM/PM")
        # Pattern: date space time with seconds
        if re.search(r':\d{2} [AP]M$', timestamp):  # Ends with ":SS AM" or ":SS PM"
            # Extract date and time parts
            parts = timestamp.split(" ", 2)
            if len(parts) == 3:
                date_part = parts[0]
                time_part = parts[1]  # "HH:MM:SS"
                am_pm = parts[2]  # "AM" or "PM"
                
                # Remove seconds from time
                time_parts = time_part.split(":")
                if len(time_parts) >= 3:  # Has seconds
                    hour_min = f"{time_parts[0]}:{time_parts[1]}"
                    # Format without leading zero on hour
                    new_timestamp = f"{date_part} {format_time_without_leading_zero(f'{hour_min} {am_pm}')}"
                    log["timestamp"] = new_timestamp
                    updated_count += 1
    
    if updated_count > 0:
        save_logs(logs)
        print(f"Updated {updated_count} log entries to remove seconds.")
    
    return updated_count > 0

def load_names():
    return load_json(NAMES_FILE, ["Alice", "Bob", "Charlie", "Diana"])

def load_rfid_cards():
    """Load RFID card mappings: {rfid_id: name}"""
    return load_json(RFID_FILE, {})

def save_rfid_cards(rfid_cards):
    """Save RFID card mappings"""
    save_json(RFID_FILE, rfid_cards)

def get_name_from_rfid(rfid_id):
    """Get the name associated with an RFID card ID"""
    rfid_cards = load_rfid_cards()
    return rfid_cards.get(rfid_id.strip(), None)

def extract_date(timestamp):
    return timestamp.split(" ")[0]

def parse_timestamp_for_sorting(timestamp):
    """
    Parse timestamp string "YYYY-MM-DD HH:MM AM/PM" into a datetime object for sorting.
    Returns a datetime object that can be used for comparison.
    """
    try:
        # Format: "2025-12-08 2:00 PM" or "2025-12-08 10:30 AM"
        # Split into date and time parts
        parts = timestamp.split(" ", 2)
        if len(parts) < 3:
            return datetime.min  # Return minimum date if parsing fails
        
        date_str = parts[0]  # "2025-12-08"
        time_str = parts[1]  # "2:00" or "10:30"
        am_pm = parts[2]     # "AM" or "PM"
        
        # Parse date
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        
        # Parse time
        time_parts = time_str.split(":")
        hour = int(time_parts[0])
        minute = int(time_parts[1]) if len(time_parts) > 1 else 0
        
        # Convert to 24-hour format
        if am_pm.upper() == "PM" and hour != 12:
            hour += 12
        elif am_pm.upper() == "AM" and hour == 12:
            hour = 0
        
        # Combine date and time
        time_obj = time(hour=hour, minute=minute)
        return datetime.combine(date_obj.date(), time_obj)
    except:
        return datetime.min  # Return minimum date if parsing fails

def get_week_range(date_str):
    """
    Get the Monday-Friday range (working week) for a given date.
    Returns (monday_date, friday_date) as strings in "YYYY-MM-DD" format.
    """
    try:
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        # Get the day of the week (0 = Monday, 6 = Sunday)
        weekday = date_obj.weekday()
        # Calculate days to subtract to get to Monday
        days_to_monday = weekday
        monday = date_obj - timedelta(days=days_to_monday)
        # Friday is 4 days after Monday
        friday = monday + timedelta(days=4)
        return monday.strftime("%Y-%m-%d"), friday.strftime("%Y-%m-%d")
    except:
        return None, None

def is_date_in_range(date_str, start_date, end_date):
    """
    Check if a date string (YYYY-MM-DD) is within the range [start_date, end_date] (inclusive).
    """
    try:
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        start_obj = datetime.strptime(start_date, "%Y-%m-%d")
        end_obj = datetime.strptime(end_date, "%Y-%m-%d")
        return start_obj <= date_obj <= end_obj
    except:
        return False

def get_state_at_timestamp(logs, timestamp):
    """
    Determine the sign-in state (IN or OUT) at a given timestamp based on logs.
    Returns the state ("IN" or "OUT") that would exist at the given timestamp.
    
    Args:
        logs: List of log entries (should be sorted chronologically)
        timestamp: The timestamp to check (format: "YYYY-MM-DD HH:MM AM/PM")
    
    Returns:
        "IN" or "OUT" - the state at that timestamp
    """
    # Sort logs chronologically
    sorted_logs = sorted(logs, key=lambda log: (parse_timestamp_for_sorting(log["timestamp"]), log["id"]))
    
    # Find the last log entry before or at the given timestamp
    target_time = parse_timestamp_for_sorting(timestamp)
    state = "OUT"  # Default state is OUT
    
    for log in sorted_logs:
        log_time = parse_timestamp_for_sorting(log["timestamp"])
        if log_time <= target_time:
            state = log["action"]  # Update state to the action at this time
        else:
            break  # We've passed the target time
    
    return state

def format_date_for_display(date_str):
    """
    Convert date string from "YYYY-MM-DD" to format like "Apr. 15" or "Jul. 5".
    Example: "2025-04-15" -> "Apr. 15"
             "2025-07-05" -> "Jul. 5"
             "2025-11-20" -> "Nov. 20"
    """
    try:
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        month_abbr = date_obj.strftime("%b")  # Apr, Jul, Nov, etc.
        day_str = date_obj.strftime("%d")  # "01", "05", "15", etc.
        day = day_str.lstrip("0") if day_str.lstrip("0") else day_str  # Remove leading zero, but keep "0" if result is empty
        return f"{month_abbr}. {day}"  # Apr. 15
    except:
        return date_str

def format_date_for_docx(date_str):
    """
    Convert date string from "YYYY-MM-DD" to format like "Nov '25" or "Mar '25".
    Example: "2025-11-20" -> "Nov '25"
             "2025-03-15" -> "Mar '25"
             "2026-06-10" -> "Jun '26"
    """
    try:
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        month_abbr = date_obj.strftime("%b")  # Nov, Mar, Jun, etc.
        year_short = date_obj.strftime("%y")  # 25, 26, etc. (without quote)
        return f"{month_abbr} '{year_short}"  # Nov '25
    except:
        return date_str

def replace_text_preserving_format(paragraph, pattern, replacement):
    """
    Replace text in a paragraph while preserving formatting (font size, bold, etc.).
    """
    if not paragraph.text:
        return False
    
    # Get the full text of the paragraph
    full_text = paragraph.text
    
    # Check if pattern exists
    if not re.search(pattern, full_text, flags=re.IGNORECASE):
        return False
    
    # Replace the text
    new_full_text = re.sub(pattern, replacement, full_text, flags=re.IGNORECASE)
    
    if new_full_text == full_text:
        return False
    
    # Get all runs to preserve formatting
    runs = list(paragraph.runs)
    if not runs:
        return False
    
    # Find the run that contains the pattern
    # Reconstruct text position by position to find which run has the match
    match = re.search(pattern, full_text, flags=re.IGNORECASE)
    if not match:
        return False
    
    match_start = match.start()
    match_end = match.end()
    
    # Find which run contains the match
    current_pos = 0
    matching_run_idx = 0
    for i, run in enumerate(runs):
        run_length = len(run.text) if run.text else 0
        if current_pos <= match_start < current_pos + run_length:
            matching_run_idx = i
            break
        current_pos += run_length
    
    # Get the formatting from the matching run (or first run if not found)
    format_run = runs[matching_run_idx] if matching_run_idx < len(runs) else runs[0]
    
    # Clear paragraph and rebuild with preserved formatting
    paragraph.clear()
    
    # Split the new text into parts: before match, replacement, after match
    before_text = new_full_text[:match_start]
    after_text = new_full_text[match_start + len(replacement):]
    
    # Add text before match (preserving first run's format)
    if before_text:
        new_run = paragraph.add_run(before_text)
        if runs:
            orig_run = runs[0]
            new_run.bold = orig_run.bold
            new_run.italic = orig_run.italic
            if orig_run.font.size:
                new_run.font.size = orig_run.font.size
            if orig_run.font.name:
                new_run.font.name = orig_run.font.name
    
    # Add replacement text (preserving matching run's format)
    new_run = paragraph.add_run(replacement)
    new_run.bold = format_run.bold
    new_run.italic = format_run.italic
    if format_run.font.size:
        new_run.font.size = format_run.font.size
    if format_run.font.name:
        new_run.font.name = format_run.font.name
    
    # Add text after match (preserving last run's format)
    if after_text:
        new_run = paragraph.add_run(after_text)
        if runs:
            orig_run = runs[-1]
            new_run.bold = orig_run.bold
            new_run.italic = orig_run.italic
            if orig_run.font.size:
                new_run.font.size = orig_run.font.size
            if orig_run.font.name:
                new_run.font.name = orig_run.font.name
    
    return True

def format_time_without_leading_zero(time_str):
    """
    Remove leading zeros from hour in time format (only the hour, not minutes).
    Example: "01:23 PM" -> "1:23 PM"
             "12:36 PM" -> "12:36 PM"
             "10:05 AM" -> "10:05 AM" (hour doesn't have leading zero, so unchanged)
    """
    # Split time into parts: "HH:MM AM/PM" or "HH:MM:SS AM/PM" (for backward compatibility)
    # Extract the time part (before AM/PM)
    time_part = time_str.split()[0] if " " in time_str else time_str
    am_pm = time_str.split()[1] if " " in time_str and len(time_str.split()) > 1 else ""
    
    parts = time_part.split(":")
    if len(parts) >= 1:
        hour_str = parts[0]
        # Remove leading zero from hour only, but keep "12" as "12" (not "2")
        if hour_str.startswith("0") and len(hour_str) == 2:
            hour = hour_str[1]  # Remove the leading zero
        else:
            hour = hour_str
        # Reconstruct with hour without leading zero (only keep hour and minute, ignore seconds)
        minute = parts[1] if len(parts) > 1 else ""
        result = f"{hour}:{minute}"
        if am_pm:
            result += f" {am_pm}"
        return result
    return time_str


# -------------------------------
# Helper for grouping logs (CSV-style)
# -------------------------------
def group_logs_csv_style(logs_raw):
    grouped = defaultdict(list)

    # IMPORTANT:
    # Sort by timestamp first (chronological order), then by ID as tiebreaker
    # This ensures logs are in chronological order even if added out of sequence
    logs_raw = sorted(logs_raw, key=lambda log: (parse_timestamp_for_sorting(log["timestamp"]), log["id"]))

    # Group by (name, date) in the order they appear
    for log in logs_raw:
        name = log["name"]
        date = extract_date(log["timestamp"])
        time = log["timestamp"].split(" ", 1)[1]
        # Format time without leading zero
        time = format_time_without_leading_zero(time)

        grouped[(name, date)].append((log["action"], time))

    final_rows = []

    # Build IN/OUT pairs in chronological order (sorted by timestamp)
    for (name, date), events in grouped.items():
        pairs = []
        current_in = None

        for action, time in events:
            if action == "IN":
                # If we already have a pending IN without OUT, create empty spot for its OUT
                if current_in is not None:
                    pairs.append((current_in, ""))
                current_in = time
            elif action == "OUT":
                if current_in:
                    # Normal case: IN followed by OUT
                    pairs.append((current_in, time))
                    current_in = None
                else:
                    # OUT without matching IN - create empty IN spot
                    pairs.append(("", time))

        # Missing OUT at the end?
        if current_in is not None:
            pairs.append((current_in, ""))

        # Split into 4-pair rows (CSV-compatible)
        for i in range(0, len(pairs), 4):
            chunk = pairs[i:i+4]
            # Format date for display (e.g., "Apr. 15" instead of "2025-04-15")
            formatted_date = format_date_for_display(date)
            final_rows.append((name, formatted_date, chunk))

    return final_rows


def export_to_docx(rows, selected_date):
    """
    Take the same rows that go into the CSV export and write them
    into the first table of the FH306 Sign-In Sheet Word document.
 
    `rows` is a list of lists in the form:
        [Name, Date, TimeIn1, TimeOut1, ..., TimeIn4, TimeOut4]
        Note: Date column now contains formatted dates like "Apr. 15"
 
    Behavior:
      - If `selected_date` is provided, write ALL rows into a single DOCX:
            FH306 Sign-In Sheet - <selected_date>.docx
      - If `selected_date` is None (exporting all dates), create ONE DOCX
        per distinct date in the data:
            FH306 Sign-In Sheet - <YYYY-MM-DD>.docx
    """
    try:
        if selected_date:
            # Single date: write everything into one document
            _write_docx_for_date(rows, selected_date)
        else:
            # All dates: group rows by date (second column) and
            # create one document per date.
            # Since dates are now formatted, we need to extract original dates from logs
            # We'll group by the formatted date string, but need original for filename
            rows_by_formatted_date = defaultdict(list)
            for row in rows:
                if len(row) < 2:
                    continue
                formatted_date = row[1]  # Formatted date like "Apr. 15"
                rows_by_formatted_date[formatted_date].append(row)
 
            # For each formatted date group, we need to find the original date
            # We'll get it from the logs by matching the formatted date
            logs = load_logs()
            for formatted_date, date_rows in rows_by_formatted_date.items():
                if not formatted_date:
                    continue
                # Find the original date format by checking logs
                # Since we can't easily reverse the format, we'll extract from a log entry
                # that matches this formatted date
                original_date = None
                for log in logs:
                    log_date = extract_date(log["timestamp"])
                    if format_date_for_display(log_date) == formatted_date:
                        original_date = log_date
                        break
                
                if original_date:
                    _write_docx_for_date(date_rows, original_date, None)
    except Exception as e:
        # Don't break the CSV export if something goes wrong
        print(f"Error while creating DOCX export: {e}")


def _write_docx_for_date(rows_for_date, date_str, output_path=None):
    """
    Helper that actually fills the template table for a single date
    and saves it as 'FH306 Sign-In Sheet - <date_str>.docx'.
    
    Args:
        rows_for_date: List of data rows to write
        date_str: Date string in "YYYY-MM-DD" format
        output_path: Optional path to save the file (if None, saves in current directory)
    """
    doc = Document(TEMPLATE_DOCX)

    # Update the date text in the document
    formatted_date = format_date_for_docx(date_str)
    # Pattern to match "FH 306 Staff and Student Sign-In ()" and replace () with the date
    # Handles variations with optional spaces: "FH 306 Staff and Student Sign-In ()" or "FH 306 Staff and Student Sign-In()"
    pattern = r'FH\s*306\s*Staff\s+and\s+Student\s+Sign-In\s*\(\)'
    replacement = f'FH 306 Staff and Student Sign-In ({formatted_date})'
    
    # Replace in all paragraphs (preserving formatting)
    for paragraph in doc.paragraphs:
        replace_text_preserving_format(paragraph, pattern, replacement)
    
    # Also check tables for date text (in case the date is in a table cell)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_preserving_format(para, pattern, replacement)
    
    # Also check headers and footers
    for section in doc.sections:
        # Check header
        if section.header:
            for paragraph in section.header.paragraphs:
                replace_text_preserving_format(paragraph, pattern, replacement)
        # Check footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                replace_text_preserving_format(paragraph, pattern, replacement)

    if not doc.tables:
        return

    table = doc.tables[0]

    # Set font size to 10 for header row (first row)
    if len(table.rows) > 0:
        header_row = table.rows[0]
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # Assume the first row is the header – keep it, clear the rest
    while len(table.rows) > 1:
        row = table.rows[-1]
        tbl = table._tbl
        tbl.remove(row._element)

    # Fill table with data rows
    for row_data in rows_for_date:
        # Ensure we always have exactly 10 columns worth of data
        row_data = (row_data + [""] * 10)[:10]

        row_cells = table.add_row().cells

        # Safely assign into available cells
        num_cells = min(len(row_cells), len(row_data))
        for i in range(num_cells):
            cell = row_cells[i]
            cell.text = str(row_data[i])
            # Set font size to 10 for all text in the cell
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    if output_path:
        out_name = output_path
    else:
        # Save to exports folder if not specified
        out_name = os.path.join(EXPORTS_DIR, DOCX_OUTPUT_FILENAME_TEMPLATE.format(date=date_str))
    doc.save(out_name)


# --- ROUTES ------------------------------------------------------------------

@app.route("/")
def index():
    names = load_names()
    logs_raw = load_logs()

    # Get today's date
    today = datetime.now().strftime("%Y-%m-%d")

    # Filter logs to only include today
    today_logs = [log for log in logs_raw if log["timestamp"].startswith(today)]

    # Group logs in CSV-style (up to 4 pairs per row)
    grouped_rows = group_logs_csv_style(today_logs)

    # Get error message from query parameters (for RFID scan errors)
    error_message = request.args.get("error")
    
    # Load RFID cards to pass to template for client-side checking
    rfid_cards = load_rfid_cards()
    
    return render_template(
        "index.html",
        names=names,
        grouped_rows=grouped_rows,
        error=error_message,
        form_name="",
        form_use_current_time=True,
        form_manual_time="",
        form_action="",
        rfid_cards=rfid_cards
    )

@app.route("/rfid_scan", methods=["POST"])
def rfid_scan():
    """Handle RFID card scan - automatically sign in/out based on current state"""
    from urllib.parse import quote
    rfid_id = request.form.get("rfid_id", "").strip()
    
    if not rfid_id:
        return redirect(f"/?error={quote('No RFID card detected')}")
    
    # Look up name from RFID card
    name = get_name_from_rfid(rfid_id)
    if not name:
        return redirect(f"/?error={quote('RFID card not registered. Please contact administrator.')}")
    
    logs = load_logs()
    names = load_names()
    
    # Get current timestamp
    now = datetime.now()
    date_str = now.strftime("%Y-%m-%d")
    time_str = now.strftime("%I:%M %p")
    timestamp = f"{date_str} {format_time_without_leading_zero(time_str)}"
    
    # Get today's logs for this person
    todays_logs = [
        log for log in logs
        if log["name"] == name and log["timestamp"].startswith(date_str)
    ]
    
    # Determine current state (at current time)
    state_at_timestamp = get_state_at_timestamp(todays_logs, timestamp)
    
    # Auto-determine action: if OUT, sign IN; if IN, sign OUT
    if state_at_timestamp == "IN":
        action = "OUT"
    else:
        action = "IN"
    
    # Validate the action
    if action == "OUT" and state_at_timestamp != "IN":
        return redirect(f"/?error={quote(f'{name} cannot sign OUT because they are not signed IN.')}")
    
    if action == "IN" and state_at_timestamp == "IN":
        return redirect(f"/?error={quote(f'{name} is already signed IN.')}")
    
    # Valid → save log
    logs.append({
        "id": len(logs),
        "name": name,
        "action": action,
        "timestamp": timestamp
    })
    
    save_logs(logs)
    return redirect("/")

@app.route("/sign", methods=["POST"])
def sign():
    name = request.form["name"].strip()
    action = request.form["action"]
    use_current_time = "use_current_time" in request.form
    manual_time = request.form.get("manual_time")

    logs = load_logs()
    names = load_names()

    # Determine timestamp
    now = datetime.now()
    date_str = now.strftime("%Y-%m-%d")

    if use_current_time:
        # Use real timestamp (format without leading zero on hour, no seconds)
        time_str = now.strftime("%I:%M %p")
        timestamp = f"{date_str} {format_time_without_leading_zero(time_str)}"
    else:
        # Use selected manual time
        # Convert 24-hour <input type="time"> into AM/PM
        # Ex: "13:45" -> 1:45 PM (no leading zero on hour, no seconds)
        if manual_time:
            manual_dt = datetime.strptime(manual_time, "%H:%M")
            time_str = manual_dt.strftime("%I:%M %p")
            timestamp = f"{date_str} {format_time_without_leading_zero(time_str)}"
        else:
            # Failsafe fallback
            time_str = now.strftime("%I:%M %p")
            timestamp = f"{date_str} {format_time_without_leading_zero(time_str)}"

    # Get today's logs for validation (for this person)
    todays_logs = [
        log for log in logs
        if log["name"] == name and log["timestamp"].startswith(date_str)
    ]

    # Determine the state at the timestamp being submitted
    # This checks what the state would be at the time being logged, not the current latest state
    state_at_timestamp = get_state_at_timestamp(todays_logs, timestamp)

    # Validation rules
    grouped_rows = group_logs_csv_style([log for log in load_logs()
                                         if log["timestamp"].startswith(date_str)])

    # Rule 1: Cannot OUT unless signed IN at that timestamp
    if action == "OUT" and state_at_timestamp != "IN":
        return render_template(
            "index.html",
            names=names,
            grouped_rows=grouped_rows,
            error=f"{name} cannot sign OUT at {timestamp.split(' ', 1)[1]} because they are not signed IN at that time.",
            form_name=name,
            form_use_current_time=use_current_time,
            form_manual_time=manual_time or "",
            form_action=action
        )

    # Rule 2: Cannot IN if already IN at that timestamp
    if action == "IN" and state_at_timestamp == "IN":
        return render_template(
            "index.html",
            names=names,
            grouped_rows=grouped_rows,
            error=f"{name} is already signed IN at {timestamp.split(' ', 1)[1]} and cannot sign in again at that time.",
            form_name=name,
            form_use_current_time=use_current_time,
            form_manual_time=manual_time or "",
            form_action=action
        )

    # Valid → save log
    logs.append({
        "id": len(logs),
        "name": name,
        "action": action,
        "timestamp": timestamp
    })

    save_logs(logs)
    return redirect("/")


# ---------- Admin & Export unchanged ----------

@app.route("/admin")
def admin():
    logs = load_logs()

    filter_type = request.args.get("filter_type", "all")
    selected_date = request.args.get("date")
    week_date = request.args.get("week_date")

    # Auto-detect filter type if not explicitly set
    if filter_type == "all":
        if selected_date:
            filter_type = "date"
        elif week_date:
            filter_type = "week"
    
    # Set default to today's date/week if filter type is set but no date/week specified
    if filter_type == "date" and not selected_date:
        selected_date = datetime.now().strftime("%Y-%m-%d")
    elif filter_type == "week" and not week_date:
        week_date = datetime.now().strftime("%Y-%m-%d")
    
    dates = sorted(set(extract_date(log["timestamp"]) for log in logs), reverse=True)

    # Filter logs based on filter type
    if filter_type == "date" and selected_date:
        logs = [log for log in logs if extract_date(log["timestamp"]) == selected_date]
    elif filter_type == "week" and week_date:
        monday, friday = get_week_range(week_date)
        if monday and friday:
            logs = [log for log in logs if is_date_in_range(extract_date(log["timestamp"]), monday, friday)]

    # Get success message if DOCX export was successful
    message = request.args.get("message", "")

    return render_template("admin.html", logs=logs, dates=dates, selected_date=selected_date, 
                          filter_type=filter_type, week_date=week_date, message=message)


@app.route("/export")
def export():
    logs = load_logs()
    selected_date = request.args.get("date")
    week_date = request.args.get("week_date")
    filename = CSV_BASE_FILENAME

    if selected_date:
        logs = [log for log in logs if extract_date(log["timestamp"]) == selected_date]
        filename = CSV_DATE_FILENAME_TEMPLATE.format(date=selected_date)
    elif week_date:
        monday, friday = get_week_range(week_date)
        if monday and friday:
            logs = [log for log in logs if is_date_in_range(extract_date(log["timestamp"]), monday, friday)]
            filename = CSV_DATE_FILENAME_TEMPLATE.format(date=f"{monday}_to_{friday}")

    # IMPORTANT: Sort by timestamp first (chronological order), then by ID as tiebreaker
    # This ensures logs are in chronological order even if added out of sequence
    logs = sorted(logs, key=lambda log: (parse_timestamp_for_sorting(log["timestamp"]), log["id"]))

    # Grouping by (name, date) preserving order
    grouped = defaultdict(list)

    for log in logs:
        name = log["name"]
        date = extract_date(log["timestamp"])
        time = log["timestamp"].split(" ", 1)[1]
        # Format time without leading zero
        time = format_time_without_leading_zero(time)

        grouped[(name, date)].append((log["action"], time))

    # We'll keep a copy of all CSV rows so we can also write them to the DOCX
    rows_for_docx = []

    with open(filename, "w", newline="") as f:
        writer = csv.writer(f)
        header = [
            "Name", "Date",
            "Time In", "Time Out",
            "Time In", "Time Out",
            "Time In", "Time Out",
            "Time In", "Time Out"
        ]
        writer.writerow(header)

        for (name, date), events in grouped.items():
            pairs = []
            current_in = None

            # Process events in chronological order (sorted by timestamp)
            for action, time in events:
                if action == "IN":
                    # If we already have a pending IN without OUT, create empty spot for its OUT
                    if current_in is not None:
                        pairs.append((current_in, ""))
                    current_in = time
                elif action == "OUT":
                    if current_in:
                        # Normal case: IN followed by OUT
                        pairs.append((current_in, time))
                        current_in = None
                    else:
                        # OUT without matching IN - create empty IN spot
                        pairs.append(("", time))

            # Missing OUT at the end?
            if current_in is not None:
                pairs.append((current_in, ""))
        
        # Split into 4-pair rows (CSV-compatible)
        for i in range(0, len(pairs), 4):
            chunk = pairs[i:i+4]
            # Format date for display (e.g., "Apr. 15" instead of "2025-04-15")
            formatted_date = format_date_for_display(date)
            row = [name, formatted_date]
            
            for pin, pout in chunk:
                row.extend([pin, pout])
            
            while len(row) < 10:
                row.extend(["", ""])
            
            writer.writerow(row)
            rows_for_docx.append(row)

    # Also write the same data into a copy of the FH306 Sign-In Sheet DOCX
    # Note: rows_for_docx has formatted dates, but export_to_docx needs original dates for grouping
    # For single date exports, pass the date. For week or all exports, pass None to create one DOCX per date
    docx_date = selected_date if selected_date else None
    export_to_docx(rows_for_docx, docx_date)

    return send_file(filename, as_attachment=True)


@app.route("/export_docx")
def export_docx():
    """
    Export DOCX files to a folder. Creates one DOCX per date, or one for selected date/week.
    """
    logs = load_logs()
    selected_date = request.args.get("date", "").strip()
    week_date = request.args.get("week_date", "").strip()
    
    # Create export folder
    if not os.path.exists(DOCX_EXPORT_FOLDER):
        os.makedirs(DOCX_EXPORT_FOLDER)
    
    # Filter logs by date or week if specified
    if selected_date:
        # Filter to only include logs for the selected date
        filtered_logs = [log for log in logs if extract_date(log["timestamp"]) == selected_date]
        if not filtered_logs:
            # No logs for selected date
            from urllib.parse import quote
            message = f"No logs found for date {selected_date}."
            return redirect(f"/admin?filter_type=date&date={quote(selected_date)}&docx_exported=0&message={quote(message)}")
        logs = filtered_logs  # Use only filtered logs
    elif week_date:
        # Filter to only include logs for the selected week
        monday, friday = get_week_range(week_date)
        if monday and friday:
            filtered_logs = [log for log in logs if is_date_in_range(extract_date(log["timestamp"]), monday, friday)]
            if not filtered_logs:
                from urllib.parse import quote
                message = f"No logs found for week {monday} to {friday}."
                return redirect(f"/admin?filter_type=week&week_date={quote(week_date)}&docx_exported=0&message={quote(message)}")
        logs = filtered_logs  # Use only filtered logs
    
    # IMPORTANT: Sort by timestamp first (chronological order), then by ID as tiebreaker
    # This ensures logs are in chronological order even if added out of sequence
    logs = sorted(logs, key=lambda log: (parse_timestamp_for_sorting(log["timestamp"]), log["id"]))
    
    # Grouping by (name, date) preserving order
    grouped = defaultdict(list)
    
    for log in logs:
        name = log["name"]
        date = extract_date(log["timestamp"])
        time = log["timestamp"].split(" ", 1)[1]
        # Format time without leading zero
        time = format_time_without_leading_zero(time)
        
        grouped[(name, date)].append((log["action"], time))
    
    # Group rows by date
    rows_by_date = defaultdict(list)
    
    for (name, date), events in grouped.items():
        pairs = []
        current_in = None
        
        # Process events in chronological order (sorted by timestamp)
        for action, time in events:
            if action == "IN":
                # If we already have a pending IN without OUT, create empty spot for its OUT
                if current_in is not None:
                    pairs.append((current_in, ""))
                current_in = time
            elif action == "OUT":
                if current_in:
                    # Normal case: IN followed by OUT
                    pairs.append((current_in, time))
                    current_in = None
                else:
                    # OUT without matching IN - create empty IN spot
                    pairs.append(("", time))
        
        # Missing OUT at the end?
        if current_in is not None:
            pairs.append((current_in, ""))
        
        # Split into 4-pair rows (CSV-compatible)
        for i in range(0, len(pairs), 4):
            chunk = pairs[i:i+4]
            # Format date for display (e.g., "Apr. 15" instead of "2025-04-15")
            formatted_date = format_date_for_display(date)
            row = [name, formatted_date]
            
            for pin, pout in chunk:
                row.extend([pin, pout])
            
            while len(row) < 10:
                row.extend(["", ""])
            
            rows_by_date[date].append(row)
    
    # Generate DOCX files
    generated_files = []
    
    # Check if a date or week is selected
    has_selected_date = selected_date and len(selected_date.strip()) > 0
    has_selected_week = week_date and len(week_date.strip()) > 0
    
    if has_selected_date:
        # Single date selected: create only one DOCX file for that date
        # Since we filtered logs above, rows_by_date should only contain the selected date
        # But we'll explicitly only process the selected date to be safe
        if selected_date in rows_by_date and len(rows_by_date[selected_date]) > 0:
            output_path = os.path.join(DOCX_EXPORT_FOLDER, DOCX_OUTPUT_FILENAME_TEMPLATE.format(date=selected_date))
            _write_docx_for_date(rows_by_date[selected_date], selected_date, output_path)
            generated_files.append(output_path)
        else:
            # This shouldn't happen if filtering worked, but handle it gracefully
            from urllib.parse import quote
            message = f"No data found for date {selected_date}."
            return redirect(f"/admin?filter_type=date&date={quote(selected_date)}&docx_exported=0&message={quote(message)}")
    elif has_selected_week:
        # Week selected: create one DOCX file per date in the week
        # Only process dates that actually have data
        monday, friday = get_week_range(week_date)
        week_dates = []
        if monday and friday:
            current = datetime.strptime(monday, "%Y-%m-%d")
            end = datetime.strptime(friday, "%Y-%m-%d")
            while current <= end:
                week_dates.append(current.strftime("%Y-%m-%d"))
                current = current + timedelta(days=1)
        
        for date in week_dates:
            if date in rows_by_date and len(rows_by_date[date]) > 0:
                output_path = os.path.join(DOCX_EXPORT_FOLDER, DOCX_OUTPUT_FILENAME_TEMPLATE.format(date=date))
                _write_docx_for_date(rows_by_date[date], date, output_path)
                generated_files.append(output_path)
    else:
        # No date/week selected: create one DOCX file per date found in rows_by_date
        # Only process dates that actually have data
        for date, date_rows in rows_by_date.items():
            if not date or len(date_rows) == 0:
                continue
            output_path = os.path.join(DOCX_EXPORT_FOLDER, DOCX_OUTPUT_FILENAME_TEMPLATE.format(date=date))
            _write_docx_for_date(date_rows, date, output_path)
            generated_files.append(output_path)
    
    # Redirect back to admin with success message
    from urllib.parse import quote
    if has_selected_date:
        message = f"Generated 1 DOCX file in '{DOCX_EXPORT_FOLDER}' folder."
        redirect_url = f"/admin?filter_type=date&date={quote(selected_date)}&docx_exported=1"
    elif has_selected_week:
        message = f"Generated {len(generated_files)} DOCX file(s) for the week in '{DOCX_EXPORT_FOLDER}' folder."
        redirect_url = f"/admin?filter_type=week&week_date={quote(week_date)}&docx_exported=1"
    else:
        message = f"Generated {len(generated_files)} DOCX files in '{DOCX_EXPORT_FOLDER}' folder."
        redirect_url = f"/admin?filter_type=all&docx_exported=1"
    
    if message:
        redirect_url += f"&message={quote(message)}"
    
    return redirect(redirect_url)


@app.route("/remove/<int:log_id>")
def remove(log_id):
    logs = load_logs()
    logs = [log for log in logs if log.get("id") != log_id]

    for i, log in enumerate(logs):
        log["id"] = i

    save_logs(logs)
    
    # Preserve filter parameters in redirect
    from urllib.parse import urlencode
    filter_type = request.args.get("filter_type", "")
    selected_date = request.args.get("date", "")
    week_date = request.args.get("week_date", "")
    
    params = {}
    if filter_type:
        params["filter_type"] = filter_type
    if selected_date:
        params["date"] = selected_date
    if week_date:
        params["week_date"] = week_date
    
    redirect_url = "/admin"
    if params:
        redirect_url += "?" + urlencode(params)
    
    return redirect(redirect_url)

@app.route("/edit/<int:log_id>", methods=["POST"])
def edit(log_id):
    logs = load_logs()
    new_ts = request.form.get("timestamp")

    # Find log
    for log in logs:
        if log["id"] == log_id:
            log["timestamp"] = new_ts
            break

    save_logs(logs)
    
    # Preserve filter parameters in redirect (from form data for POST requests)
    from urllib.parse import urlencode
    filter_type = request.form.get("filter_type", "")
    selected_date = request.form.get("date", "")
    week_date = request.form.get("week_date", "")
    
    params = {}
    if filter_type:
        params["filter_type"] = filter_type
    if selected_date:
        params["date"] = selected_date
    if week_date:
        params["week_date"] = week_date
    
    redirect_url = "/admin"
    if params:
        redirect_url += "?" + urlencode(params)
    
    return redirect(redirect_url)

@app.route("/rfid")
def rfid_management():
    """RFID card management page"""
    names = load_names()
    rfid_cards = load_rfid_cards()
    message = request.args.get("message", "")
    
    return render_template("rfid.html", names=names, rfid_cards=rfid_cards, message=message)

@app.route("/rfid/add", methods=["POST"])
def add_rfid_card():
    """Add or update RFID card association"""
    from urllib.parse import quote
    name = request.form.get("name", "").strip()
    rfid_id = request.form.get("rfid_id", "").strip()
    
    if not name or not rfid_id:
        return redirect(f"/rfid?error={quote('Name and RFID ID are required')}")
    
    rfid_cards = load_rfid_cards()
    rfid_cards[rfid_id] = name
    save_rfid_cards(rfid_cards)
    
    return redirect("/rfid?message=RFID card linked successfully")

@app.route("/rfid/remove/<rfid_id>")
def remove_rfid_card(rfid_id):
    """Remove RFID card association"""
    from urllib.parse import quote
    rfid_cards = load_rfid_cards()
    if rfid_id in rfid_cards:
        del rfid_cards[rfid_id]
        save_rfid_cards(rfid_cards)
        return redirect("/rfid?message=RFID card removed successfully")
    return redirect(f"/rfid?error={quote('RFID card not found')}")


if __name__ == "__main__":
    # Update existing logs to remove seconds on startup (only runs once)
    remove_seconds_from_logs()
    app.run(debug=True, port=8080)
